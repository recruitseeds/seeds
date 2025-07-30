# text_extractor.py
import io
import logging

import docx
import fitz  # PyMuPDF
from fastapi import HTTPException, status

logger = logging.getLogger(__name__)

SUPPORTED_CONTENT_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}


def get_file_extension(file_key: str) -> str | None:
    if '.' in file_key:
        base_key = file_key.split('?')[0].split('#')[0]
        if '.' in base_key:
            return base_key.split('.')[-1].lower()
    return None


def extract_text_and_links_from_pdf(file_content: bytes) -> tuple[str, list[str]]:
    logger.info("--- EXTRACTOR: ENTERED extract_text_and_links_from_pdf ---")
    text_parts = []  # Use a list to join later
    links = []
    try:
        logger.info(
            "--- EXTRACTOR: Attempting to open PDF with fitz.open()... ---")
        with fitz.open(stream=file_content, filetype="pdf") as doc:
            logger.info(
                f"--- EXTRACTOR: PDF opened. Processing {len(doc)} pages. ---")
            for page_num, page in enumerate(doc):
                logger.info(
                    f"--- EXTRACTOR: Processing PDF page {page_num + 1}... ---")
                page_text = page.get_text("text", sort=True)
                if page_text:
                    text_parts.append(page_text)  # Add page text to list

                page_links = page.get_links()
                if page_links:
                    for link in page_links:
                        if link.get('kind') == fitz.LINK_URI and link.get('uri'):
                            uri = link.get('uri')
                            if isinstance(uri, str) and uri.strip():
                                links.append(uri.strip())
            logger.info(
                "--- EXTRACTOR: Finished iterating through PDF pages. ---")
    except Exception as e:
        logger.error(
            f"--- EXTRACTOR: Error extracting text/links from PDF: {e} ---", exc_info=True)

    text = "\n".join(text_parts)  # Join pages with newlines
    unique_links = sorted(list(set(links)))
    logger.info(
        f"--- EXTRACTOR: Extracted {len(text)} chars and {len(unique_links)} unique links from PDF. ---")
    logger.info(
        "--- EXTRACTOR: EXITING extract_text_and_links_from_pdf NORMALLY ---")
    return text, unique_links


def extract_text_from_docx(file_content: bytes) -> str:
    logger.info("--- EXTRACTOR: ENTERED extract_text_from_docx ---")
    text_parts = []  # Use a list to join paragraphs
    try:
        logger.info(
            "--- EXTRACTOR: Attempting to open DOCX with docx.Document()... ---")
        doc = docx.Document(io.BytesIO(file_content))
        logger.info(
            "--- EXTRACTOR: DOCX document opened. Iterating paragraphs... ---")
        for para_num, para in enumerate(doc.paragraphs):
            text_parts.append(para.text)  # Add each paragraph's text
        logger.info(
            "--- EXTRACTOR: Finished iterating through DOCX paragraphs. ---")
    except Exception as e:
        logger.error(
            f"--- EXTRACTOR: Error extracting text from DOCX: {e} ---", exc_info=True)

    text = "\n".join(text_parts)  # Join paragraphs with newline characters
    logger.info(
        f"--- EXTRACTOR: Extracted {len(text)} characters from DOCX. ---")
    logger.info("--- EXTRACTOR: EXITING extract_text_from_docx NORMALLY ---")
    return text


def extract_text(file_content: bytes, file_key: str) -> tuple[str, list[str]]:
    logger.info(f"--- EXTRACTOR: ENTERED extract_text for key: {file_key} ---")
    extension = get_file_extension(file_key)
    logger.info(
        f"--- EXTRACTOR: Attempting text extraction for file key: {file_key} (extension: {extension}) ---")
    text = ""
    annotation_links = []

    if extension == "pdf":
        logger.info("--- EXTRACTOR: Processing as PDF... ---")
        text, annotation_links = extract_text_and_links_from_pdf(file_content)
        logger.info("--- EXTRACTOR: Finished processing PDF. ---")
    elif extension == "docx":
        logger.info("--- EXTRACTOR: Processing as DOCX... ---")
        # annotation_links will be empty for docx by default
        text = extract_text_from_docx(file_content)
        logger.info("--- EXTRACTOR: Finished processing DOCX. ---")
    else:
        logger.warning(
            f"--- EXTRACTOR: Unsupported file type received: {extension} for key: {file_key} ---")
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"Unsupported file type: '{extension}'. Only PDF and DOCX are currently supported."
        )

    if not text and not annotation_links:  # For PDF, text could be empty if only images
        logger.warning(
            f"--- EXTRACTOR: Could not extract text or links (text empty: {not text}, links empty: {not annotation_links}) from file: {file_key} ---")

    logger.info(
        f"--- EXTRACTOR: EXITING extract_text NORMALLY for key: {file_key} ---")
    return text, annotation_links
